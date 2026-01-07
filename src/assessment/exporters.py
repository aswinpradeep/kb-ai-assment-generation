import json
import logging
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

# WeasyPrint Import
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

logger = logging.getLogger(__name__)

# Suppress noisy logs from pdf generation libraries
# WeasyPrint and FontTools can be very verbose, especially during font subsetting
# The user specifically requested only error logs.
for logger_name in ["weasyprint", "fontTools", "fontTools.subset", "fontTools.ttLib", "pydyf"]:
    logging.getLogger(logger_name).setLevel(logging.ERROR)

RESOURCE_DIR = Path(__file__).parent / "resources" / "fonts"

def get_css_font_faces() -> str:
    """Generates CSS @font-face rules for all available Noto fonts."""
    font_map = {
        "NotoSansDevanagari-Regular.ttf": "NotoSansDevanagari",
        "NotoSansTamil-Regular.ttf": "NotoSansTamil",
        "NotoSansTelugu-Regular.ttf": "NotoSansTelugu",
        "NotoSansKannada-Regular.ttf": "NotoSansKannada",
        "NotoSansMalayalam-Regular.ttf": "NotoSansMalayalam",
        "NotoSansBengali-Regular.ttf": "NotoSansBengali",
        "NotoSansGujarati-Regular.ttf": "NotoSansGujarati",
        "NotoSansGurmukhi-Regular.ttf": "NotoSansGurmukhi"
    }
    
    css = []
    for filename, font_family in font_map.items():
        font_path = RESOURCE_DIR / filename
        if font_path.exists():
            # WeasyPrint needs file:// URI for local files or absolute paths
            css.append(f"""
            @font-face {{
                font-family: '{font_family}';
                src: url('file://{font_path.absolute()}');
            }}""")
            
    return "\n".join(css)

def generate_html_content(assessment_data: dict) -> str:
    """Constructs the HTML report string."""
    blueprint = assessment_data.get("blueprint", {})
    questions_obj = assessment_data.get("questions", {})
    
    # Audit Data
    prompt_ver = blueprint.get("prompt_version", "N/A")
    api_ver = blueprint.get("api_version", "N/A")
    scope = blueprint.get('assessment_scope_summary', 'N/A')
    
    # CSS
    font_faces = get_css_font_faces()
    # Pango Stack: Put specific fonts first, then sans-serif fallback
    font_stack = "'NotoSansMalayalam', 'NotoSansDevanagari', 'NotoSansTamil', 'NotoSansTelugu', 'NotoSansKannada', 'NotoSansBengali', 'NotoSansGujarati', 'NotoSansGurmukhi', sans-serif"

    html_parts = ["""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            %s
            body {
                font-family: %s;
                font-size: 11pt;
                line-height: 1.5;
                color: #333;
                margin: 40px;
            }
            h1 { font-size: 24pt; color: #2c3e50; border-bottom: 2px solid #eee; padding-bottom: 10px; }
            h2 { font-size: 18pt; color: #34495e; margin-top: 30px; }
            h3 { font-size: 14pt; color: #7f8c8d; }
            .audit-table { width: 100%%; border-collapse: collapse; margin-bottom: 20px; }
            .audit-table th, .audit-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .audit-table th { background-color: #f2f2f2; }
            .question-block { margin-bottom: 25px; page-break-inside: avoid; }
            .question-text { font-weight: bold; font-size: 12pt; margin-bottom: 8px; }
            .options-list { list-style-type: none; padding-left: 20px; }
            .options-list li { margin-bottom: 4px; }
            .reasoning-box { background-color: #f8f9fa; border-left: 4px solid #3498db; padding: 10px; margin-top: 10px; font-size: 10pt; }
            .correct { color: #27ae60; font-weight: bold; }
        </style>
    </head>
    <body>
        <h1>Course Assessment Report</h1>
        
        <p><b>Assessment Scope:</b> %s</p>
        
        <h3>Audit Information</h3>
        <table class="audit-table">
            <tr><th>Field</th><th>Value</th></tr>
            <tr><td>Prompt Version</td><td>%s</td></tr>
            <tr><td>API Version</td><td>%s</td></tr>
        </table>
        
        <h2>Questions & Reasoning</h2>
    """ % (font_faces, font_stack, scope, prompt_ver, api_ver)]

    # Dynamic Questions
    q_counter = 1
    for q_type, q_list in questions_obj.items():
        html_parts.append(f"<h3>{q_type} ({len(q_list)})</h3>")
        
        for q in q_list:
            q_txt = q.get("question_text", "N/A")
            q_html = f"""
            <div class="question-block">
                <div class="question-text">Q{q_counter}: {q_txt}</div>
            """
            
            # Options / Body
            if q_type == "Multiple Choice Question":
                opts_html = "".join([f"<li>- {o.get('text', '')}</li>" for o in q.get("options", [])])
                idx = q.get('correct_option_index')
                q_html += f"<ul class='options-list'>{opts_html}</ul>"
                q_html += f"<div class='correct'>Correct Answer: Option {idx}</div>"
            
            elif q_type == "MTF Question":
                pairs_html = "".join([f"<li>- {p.get('left')} &rarr; {p.get('right')}</li>" for p in q.get("pairs", [])])
                q_html += f"<ul class='options-list'>{pairs_html}</ul>"
            
            elif q_type == "Multi-Choice Question":
                opts_html = "".join([f"<li>[ ] {o.get('text', '')}</li>" for o in q.get("options", [])])
                corr = q.get('correct_option_index')
                corr_str = ", ".join(map(str, corr)) if isinstance(corr, list) else str(corr)
                q_html += f"<ul class='options-list'>{opts_html}</ul>"
                q_html += f"<div class='correct'>Correct Options: {corr_str}</div>"
            
            elif q_type == "True/False Question":
                q_html += "<ul class='options-list'><li>- True</li><li>- False</li></ul>"
                q_html += f"<div class='correct'>Correct Answer: {q.get('correct_answer')}</div>"
            
            else:
                q_html += f"<div class='correct'>Answer: {q.get('correct_answer')}</div>"

            # Reasoning Box
            rs = q.get("reasoning", {})
            kcm = rs.get("competency_alignment", {}).get("kcm", {})
            
            q_html += f"""
                <div class="reasoning-box">
                    <b>Rationale:</b> {rs.get('question_type_rationale')}<br/>
                    <b>Bloom's Level:</b> {q.get('blooms_level')} ({rs.get('blooms_level_justification')})<br/>
                    <b>Competency:</b> {kcm.get('competency_area')} - {kcm.get('competency_theme')}<br/>
                    <b>Relevance:</b> {q.get('relevance_percentage')}%
                </div>
            </div>
            """
            html_parts.append(q_html)
            q_counter += 1

    html_parts.append("</body></html>")
    return "\n".join(html_parts)

def generate_pdf(assessment_data: dict, output_path: Path):
    """Generates a PDF report using WeasyPrint (HTML-to-PDF)."""
    try:
        html_content = generate_html_content(assessment_data)
        font_config = FontConfiguration()
        HTML(string=html_content).write_pdf(target=str(output_path), font_config=font_config)
        logger.info(f"Generated PDF with WeasyPrint: {output_path}")
    except Exception as e:
        logger.error(f"WeasyPrint PDF Generation Failed: {e}")
        # Fallback? No, fail explicitly is better than bad boxes.
        raise


def generate_docx(assessment_data: dict, output_path: Path):
    """Generates a DOCX report from the assessment JSON data."""
    doc = Document()
    doc.add_heading('Course Assessment Report', 0)

    blueprint = assessment_data.get("blueprint", {})
    doc.add_paragraph(f"Assessment Scope: {blueprint.get('assessment_scope_summary', 'N/A')}")
    
    # Audit Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Prompt Version'
    row_cells[1].text = str(blueprint.get("prompt_version", "N/A"))
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'API Version'
    row_cells[1].text = str(blueprint.get("api_version", "N/A"))
    
    doc.add_heading('Questions & Reasoning', level=1)
    
    questions_obj = assessment_data.get("questions", {})
    
    for q_type, q_list in questions_obj.items():
        doc.add_heading(f"{q_type} ({len(q_list)})", level=2)
        
        for i, q in enumerate(q_list, 1):
            doc.add_paragraph(f"Q{i}: {q.get('question_text', 'N/A')}", style='List Number')
            
            if q_type == "Multiple Choice Question":
                for opt in q.get("options", []):
                    doc.add_paragraph(f"- {opt.get('text', '')}", style='List Bullet')
                p = doc.add_paragraph()
                p.add_run(f"Correct Answer: Option {q.get('correct_option_index')}").bold = True
            
            elif q_type == "MTF Question":
                for p_item in q.get("pairs", []):
                    doc.add_paragraph(f"- {p_item.get('left')} -> {p_item.get('right')}", style='List Bullet')

            elif q_type == "Multi-Choice Question":
                for opt in q.get("options", []):
                     doc.add_paragraph(f"[ ] {opt.get('text', '')}", style='List Bullet')
                p = doc.add_paragraph()
                corr = q.get('correct_option_index')
                corr_str = ", ".join(map(str, corr)) if isinstance(corr, list) else str(corr)
                p.add_run(f"Correct Options: {corr_str}").bold = True

            elif q_type == "True/False Question":
                 doc.add_paragraph(f"- True", style='List Bullet')
                 doc.add_paragraph(f"- False", style='List Bullet')
                 p = doc.add_paragraph()
                 p.add_run(f"Correct Answer: {q.get('correct_answer')}").bold = True
            
            else:
                 p = doc.add_paragraph()
                 p.add_run(f"Answer: {q.get('correct_answer')}").bold = True

            # Reasoning
            reasoning = q.get("reasoning", {})
            kcm = reasoning.get("competency_alignment", {}).get("kcm", {})
            
            r_para = doc.add_paragraph()
            r_para.add_run(f"\nRationale: {reasoning.get('question_type_rationale')}\n").italic = True
            r_para.add_run(f"Bloom's: {q.get('blooms_level')} | Relevance: {q.get('relevance_percentage')}%\n")
            r_para.add_run(f"Competency: {kcm.get('competency_area')} - {kcm.get('competency_theme')}")
            
    doc.save(str(output_path))
